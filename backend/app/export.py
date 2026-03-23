"""Export medical documents to PDF and DOCX formats."""

import io
import re
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from .config import EXPORTS_DIR

logger = logging.getLogger(__name__)


def _register_fonts():
    """Register a font that supports Cyrillic."""
    # Try to find a system font with Cyrillic support
    font_paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
    ]
    bold_paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
    ]
    
    font_registered = False
    for fp in font_paths:
        if Path(fp).exists():
            pdfmetrics.registerFont(TTFont("CyrillicFont", fp))
            font_registered = True
            break

    for fp in bold_paths:
        if Path(fp).exists():
            pdfmetrics.registerFont(TTFont("CyrillicFontBold", fp))
            break
    
    if not font_registered:
        logger.warning("No Cyrillic font found, PDF may not render correctly")
        return "Helvetica", "Helvetica-Bold"
    
    return "CyrillicFont", "CyrillicFontBold"


def _parse_markdown_sections(markdown_text: str):
    """Parse markdown into structured sections."""
    sections = []
    current_section = {"title": "", "content": []}
    
    for line in markdown_text.split("\n"):
        stripped = line.strip()
        
        if stripped.startswith("## "):
            if current_section["title"] or current_section["content"]:
                sections.append(current_section)
            current_section = {"title": stripped[3:].strip(), "content": []}
        elif stripped:
            # Clean markdown formatting
            clean = stripped
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)  # bold
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)        # italic
            current_section["content"].append(clean)
    
    if current_section["title"] or current_section["content"]:
        sections.append(current_section)
    
    return sections


def export_to_docx(markdown_text: str, patient_info: str = "") -> bytes:
    """Export structured medical text to DOCX format."""
    doc = Document()
    
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    
    # Title
    title = doc.add_heading("Медицинская запись", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta info
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}").font.size = Pt(10)
    if patient_info:
        meta.add_run(f"\n{patient_info}").font.size = Pt(10)
    
    doc.add_paragraph()  # spacer
    
    # Parse and add sections
    sections = _parse_markdown_sections(markdown_text)
    
    for section in sections:
        if section["title"]:
            heading = doc.add_heading(section["title"], level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 51, 102)
        
        for line in section["content"]:
            p = doc.add_paragraph()
            if line.startswith("- ") or line.startswith("• "):
                p.style = "List Bullet"
                p.add_run(line[2:])
            elif re.match(r'^\d+[\.\)]\s', line):
                p.style = "List Number"
                p.add_run(re.sub(r'^\d+[\.\)]\s', '', line))
            else:
                p.add_run(line)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Документ сформирован автоматически системой МедЗапись AI")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_pdf(markdown_text: str, patient_info: str = "") -> bytes:
    """Export structured medical text to PDF format."""
    font_name, font_bold = _register_fonts()
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles with Cyrillic font
    title_style = ParagraphStyle(
        "MedTitle",
        parent=styles["Heading1"],
        fontName=font_bold,
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    
    heading_style = ParagraphStyle(
        "MedHeading",
        parent=styles["Heading2"],
        fontName=font_bold,
        fontSize=13,
        textColor=RGBColor(0, 51, 102),  # reportlab can use this
        spaceBefore=12,
        spaceAfter=6,
    )
    # Fix color for reportlab
    heading_style.textColor = "#003366"
    
    body_style = ParagraphStyle(
        "MedBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=14,
        spaceAfter=3,
    )
    
    meta_style = ParagraphStyle(
        "MedMeta",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9,
        alignment=TA_CENTER,
        textColor="#888888",
    )
    
    bullet_style = ParagraphStyle(
        "MedBullet",
        parent=body_style,
        bulletIndent=10,
        leftIndent=25,
    )
    
    elements = []
    
    # Title
    elements.append(Paragraph("Медицинская запись", title_style))
    date_str = datetime.now().strftime('%d.%m.%Y %H:%M')
    elements.append(Paragraph(f"Дата: {date_str}  {patient_info}", meta_style))
    elements.append(Spacer(1, 10))
    
    # Sections
    sections = _parse_markdown_sections(markdown_text)
    for section in sections:
        if section["title"]:
            elements.append(Paragraph(section["title"], heading_style))
        
        for line in section["content"]:
            if line.startswith("- ") or line.startswith("• "):
                elements.append(Paragraph(f"• {line[2:]}", bullet_style))
            elif re.match(r'^\d+[\.\)]\s', line):
                elements.append(Paragraph(line, bullet_style))
            else:
                elements.append(Paragraph(line, body_style))
    
    # Footer
    elements.append(Spacer(1, 20))
    elements.append(Paragraph(
        "Документ сформирован автоматически системой МедЗапись AI",
        meta_style,
    ))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()
